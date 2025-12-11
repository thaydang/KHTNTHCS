"""
exam_to_docx.py
Converts Markdown exam files to professionally formatted .docx files
Compatible with GitHub Actions for automated document generation
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import sys
from pathlib import Path
try:
    from latex2mathml.converter import convert
except ImportError:
    convert = None
from io import BytesIO
import matplotlib.pyplot as plt

class ExamDocumentGenerator:
    """Generate professional exam documents from Markdown"""

    def __init__(self, md_file_path, output_dir=None):
        self.md_file = md_file_path
        self.output_dir = Path(output_dir) if output_dir else Path(md_file_path).parent
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """Configure document styles"""
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

    def add_heading(self, text, level=1):
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return heading

    def add_equation_to_paragraph(self, paragraph, latex_str):
        """Render LaTeX equation to an image and insert into the paragraph.

        Note: this embeds the equation as a PNG image (visually identical) since
        creating native editable OMML equations is non-trivial.
        """
        latex = latex_str.strip()
        # Ensure math delimiters for matplotlib
        if not (latex.startswith('$') and latex.endswith('$')):
            math_text = f"${latex}$"
        else:
            math_text = latex

        try:
            # Render using matplotlib to a PNG in-memory
            fig = plt.figure()
            # invisible figure
            fig.patch.set_alpha(0.0)
            txt = fig.text(0, 0, math_text, fontsize=14)
            buffer = BytesIO()
            fig.savefig(buffer, dpi=200, bbox_inches='tight', pad_inches=0.02, transparent=True)
            plt.close(fig)
            buffer.seek(0)

            run = paragraph.add_run()
            run.add_picture(buffer, width=Inches(1.6))
            return run
        except Exception as e:
            # Fallback: insert LaTeX source as plain text
            print(f"Error rendering equation '{latex}': {e}")
            return paragraph.add_run(latex)

    def add_info_table(self):
        table = self.doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = "Nội dung"
        table.rows[0].cells[1].text = "Chi tiết"
        table.rows[1].cells[0].text = "Trường"
        table.rows[1].cells[1].text = "................................................................."
        table.rows[2].cells[0].text = "Lớp"
        table.rows[2].cells[1].text = "................................................................."
        self.doc.add_paragraph()

    def add_metadata(self, metadata):
        for k, v in metadata.items():
            p = self.doc.add_paragraph()
            run = p.add_run(f"{k}: ")
            run.bold = True
            p.add_run(v)
        self.doc.add_paragraph()

    def add_paragraphs_from_text(self, text):
        # Convert LaTeX equations in the text
        lines = text.splitlines()
        i = 0
        buffer = []
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            if line.startswith('# '):
                if buffer:
                    self.add_paragraph_with_equations('\n'.join(buffer))
                    buffer = []
                self.add_heading(line.lstrip('# ').strip(), level=1)
            elif line.startswith('## '):
                if buffer:
                    self.add_paragraph_with_equations('\n'.join(buffer))
                    buffer = []
                self.add_heading(line.lstrip('## ').strip(), level=2)
            else:
                buffer.append(line)
            i += 1
        
        if buffer:
            self.add_paragraph_with_equations('\n'.join(buffer))

    def add_paragraph_with_equations(self, text):
        """Add paragraph with proper handling of LaTeX equations"""
        # Split by $ to find equations
        parts = re.split(r'(\$[^$]+\$)', text)
        
        if len(parts) == 1:
            # No equations, just add normal paragraph
            if text.strip():
                self.doc.add_paragraph(text)
        else:
            # Mix of text and equations
            p = self.doc.add_paragraph()
            for part in parts:
                if part.startswith('$') and part.endswith('$'):
                    # This is an equation
                    latex = part[1:-1]  # Remove $ delimiters
                    self.add_equation_to_paragraph(p, latex)
                elif part:
                    # Regular text
                    p.add_run(part)

    def save(self, output_filename=None):
        if output_filename:
            outpath = self.output_dir / output_filename
        else:
            base_name = Path(self.md_file).stem
            outpath = self.output_dir / f"{base_name}-Formatted.docx"
        self.doc.save(str(outpath))
        print(f"Document saved: {outpath}")
        return outpath

    def parse_and_generate(self):
        text = Path(self.md_file).read_text(encoding='utf-8')
        # Try to find a main title
        m = re.search(r'^\s*#\s*(.+)$', text, flags=re.MULTILINE)
        if m:
            self.add_heading(m.group(1).strip(), level=1)
        # Add info table and metadata if present
        metadata = {}
        mm = re.search(r'\*\*Môn học:\*\*\s*(.+)', text)
        if mm:
            metadata['Môn học'] = mm.group(1).strip()
        mm = re.search(r'\*\*Khối lớp:\*\*\s*(.+)', text)
        if mm:
            metadata['Khối lớp'] = mm.group(1).strip()
        mm = re.search(r'\*\*Thời gian làm bài:\*\*\s*(.+)', text)
        if mm:
            metadata['Thời gian'] = mm.group(1).strip()
        mm = re.search(r'\*\*Phạm vi kiến thức:\*\*\s*(.+)', text)
        if mm:
            metadata['Phạm vi'] = mm.group(1).strip()
        self.add_info_table()
        if metadata:
            self.add_metadata(metadata)
        # Add remaining content naively
        self.add_paragraphs_from_text(text)

def batch_convert(directory, output_dir=None):
    directory = Path(directory)
    md_files = list(directory.rglob("*.md"))
    if not md_files:
        print("No markdown files found.")
        return
    outdir = Path(output_dir) if output_dir else directory / "output"
    outdir.mkdir(parents=True, exist_ok=True)
    for md in md_files:
        try:
            print("Processing", md)
            gen = ExamDocumentGenerator(str(md), outdir)
            gen.parse_and_generate()
            gen.save(md.stem + ".docx")
        except Exception as e:
            print("Error:", e)

def main():
    if len(sys.argv) >= 2 and sys.argv[1] == "--batch":
        directory = sys.argv[2] if len(sys.argv) > 2 else "."
        output_dir = sys.argv[3] if len(sys.argv) > 3 else None
        batch_convert(directory, output_dir)
    elif len(sys.argv) >= 2:
        md = sys.argv[1]
        output_filename = sys.argv[2] if len(sys.argv) > 2 else None
        gen = ExamDocumentGenerator(md, Path(md).parent)
        gen.parse_and_generate()
        gen.save(output_filename)
    else:
        print("Usage:")
        print("  python exam_to_docx.py <markdown_file> [output_dir]")
        print("  python exam_to_docx.py --batch <directory> [output_dir]")

if __name__ == '__main__':
    main()
